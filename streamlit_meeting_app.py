import streamlit as st
import os
from openai import OpenAI
from docx import Document
from io import BytesIO

api_key = st.secrets["OPENAI_API_KEY"]

# Initialize OpenAI client
client = OpenAI(api_key=api_key)

# Function to transcribe audio using Whisper
def transcribe_audio(audio_file):
    transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file)
    return transcription['text'] if isinstance(transcription, dict) else transcription.text

# Function to generate response based on prompt and model
def generate_response(transcription, model, custom_prompt):
    response = client.chat.completions.create(
        model=model,
        temperature=0,
        messages=[
            {"role": "system", "content": custom_prompt},
            {"role": "user", "content": transcription}
        ]
    )
    return response.choices[0].message.content

# Function to save meeting minutes as a Word document
def save_as_docx(minutes):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to convert video files to .mp3
def convert_video_to_mp3(uploaded_file, suffix):
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_video_file:
        temp_video_file.write(uploaded_file.getbuffer())
        temp_video_file_path = temp_video_file.name

    video = mp.VideoFileClip(temp_video_file_path)

    if video.audio is None:
        st.error(f"The uploaded {suffix} file does not contain an audio track.")
        return None

    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as audio_file:
        audio_file_path = audio_file.name

    video.audio.write_audiofile(audio_file_path)
    return audio_file_path

# Function to read text from a .docx file
def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# Function to read text from a .txt file
def read_txt(file):
    return file.read().decode("utf-8")

# Pre-canned prompts and their respective headings
pre_canned_prompts = {
    "meeting_summary": {
        "summary": {
            "prompt": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.",
            "heading": "Summary"
        },
        "key_points": {
            "prompt": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.",
            "heading": "Key Points"
        },
        "action_items": {
            "prompt": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely.",
            "heading": "Action Items"
        },
        "sentiment": {
            "prompt": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible.",
            "heading": "Sentiment Analysis"
        }
    },
    "user_research": {
        "summary": {
            "prompt": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points.",
            "heading": "Summary"
        },
        "biographical_info": {
            "prompt": "You are a proficient AI with a specialty in distilling biographical information about people. Based on the following text, please identify biographical information about the subject of the research study.",
            "heading": "Biographical Info"
        },
        "key_insights": {
            "prompt": "You are a proficient AI with a specialty in distilling information into key points. Based on the following user research transcript, please identify the key insights. Identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about.",
            "heading": "Key Insights"
        },
        "recommendations": {
            "prompt": "You are a proficient AI with a specialty in identifying meaningful product opportunities. Based on the transcript, please identify product recommendations/opportunities.",
            "heading": "Recommendations"
        }
    }
}

# Streamlit app
def main():
    st.title("Meeting Summarizer")

    st.info("Upload an mp3, mp4, mov, docx, or txt file to start!")
    uploaded_file = st.file_uploader("Upload an audio, video, or text file", type=["mp3", "mp4", "mov", "docx", "txt"])
    if uploaded_file is not None:
        transcription = None

        if uploaded_file.type in ["video/quicktime", "video/mp4"]:
            suffix = ".mov" if uploaded_file.type == "video/quicktime" else ".mp4"
            audio_file_path = convert_video_to_mp3(uploaded_file, suffix)
            if audio_file_path is not None:
                with open(audio_file_path, "rb") as f:
                    transcription = transcribe_audio(f)
        elif uploaded_file.type == "audio/mpeg":
            transcription = transcribe_audio(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            transcription = read_docx(uploaded_file)
        elif uploaded_file.type == "text/plain":
            transcription = read_txt(uploaded_file)

        if transcription:
            st.subheader("Transcription")
            st.write(transcription)

            st.info("Select what you'd like to create!")
            summary_type = st.radio(
                "Select the type of summary you want to generate:",
                ("", "Meeting Summary", "User Research Synthesis"),
                index=0
            )

            if 'prompts' not in st.session_state:
                st.session_state.prompts = []

            checkboxes = {}
            if summary_type == "Meeting Summary":
                st.markdown("### Meeting Summary Prompts")
                st.info("Select the sections you'd like in your document!")
                checkboxes = {
                    "summary": st.checkbox("Summary"),
                    "key_points": st.checkbox("Key Points"),
                    "action_items": st.checkbox("Action Items"),
                    "sentiment": st.checkbox("Sentiment Analysis")
                }

            elif summary_type == "User Research Synthesis":
                st.markdown("### User Research Synthesis Prompts")
                st.info("Select the sections you'd like in your document!")
                checkboxes = {
                    "summary": st.checkbox("Summary", key="user_summary"),
                    "biographical_info": st.checkbox("Biographical Info"),
                    "key_insights": st.checkbox("Key Insights"),
                    "recommendations": st.checkbox("Recommendations")
                }

            if any(checkboxes.values()):
                st.info("Click 'Create GPT Tasks' to proceed")
                if st.button("Create GPT Tasks"):
                    for key, checked in checkboxes.items():
                        if checked:
                            st.session_state.prompts.append({
                                "prompt": pre_canned_prompts[summary_type.lower().replace(" ", "_")][key]["prompt"],
                                "model": "gpt-4o",
                                "heading": pre_canned_prompts[summary_type.lower().replace(" ", "_")][key]["heading"]
                            })

            for i, prompt_info in enumerate(st.session_state.prompts):
                st.subheader(f"GPT Task {i+1} - {prompt_info['heading']}")
                st.info("Update the pre-canned prompt to customize!")
                prompt_info["model"] = st.text_input("Model", value=prompt_info["model"], key=f"model_{i}")
                prompt_info["prompt"] = st.text_area("Prompt", value=prompt_info["prompt"], key=f"prompt_{i}")
                if st.button("Remove GPT Task", key=f"remove_gpt_task_{i}"):
                    st.session_state.prompts.pop(i)
                    break

            if st.session_state.prompts:
                st.info("Click generate to create your document!")
                st.markdown(
                    """
                    <style>
                    .blue-button button {
                        background-color: #007BFF !important;
                        color: white !important;
                    }
                    </style>
                    """,
                    unsafe_allow_html=True
                )
                if st.button("Generate", key="generate", help=None, on_click=None, disabled=False, use_container_width=False):
                    minutes = {}
                    for i, prompt_info in enumerate(st.session_state.prompts):
                        task_key = prompt_info["heading"] if prompt_info["heading"] else f"Task {i+1}"
                        minutes[task_key] = generate_response(transcription, prompt_info["model"], prompt_info["prompt"])
                    st.session_state.generated_minutes = minutes  # Store the generated minutes in session state

            # Display generated minutes if they exist in session state
            if 'generated_minutes' in st.session_state:
                st.subheader("Meeting Minutes")
                for key, value in st.session_state.generated_minutes.items():
                    st.write(f"**{key}**")
                    st.write(value)

                docx_file = save_as_docx(st.session_state.generated_minutes)

                st.info("Click download to get a docx file of your document!")
                st.download_button(
                    label="Download Meeting Minutes",
                    data=docx_file,
                    file_name="meeting_minutes.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()
